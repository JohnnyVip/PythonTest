import docx
doc=docx.Document("a.docx")
content="\n".join([para.text for para in doc.paragraphs])
a=content.split("\n")
b=[]
for i in a:
    if "strings you want to search" in i:
        b.append(i)
####print(b)####b中包含了所有的查找字符串的段落啦。复制到word中就完了。
file=docx.Document()
for i in b:
    file.add_paragraph(i)
file.save("b.docx")